import streamlit as st
from docx import Document
from PIL import Image
import io
import pandas as pd

def remove_duplicates(input_list : list) -> list:
    seen = set()
    result = []
    for item in input_list:
        if item not in seen:
            seen.add(item)
            result.append(item)
    return result

def tranform(arr : list):
    a = [[arr[0][0]],[arr[0][1]]]
    tmp1 = []
    tmp2 = []
    for i in range(1,len(arr)):
        tmp1.append(arr[i][0])
        tmp2.append(arr[i][1])
    tmp2 = remove_duplicates(tmp2)
    t = len(tmp2)

    while(t < len(tmp1)):
        tmp2.append("")
        t += 1
    
    a[0].append(tmp1)
    a[1].append(tmp2)
    a = dict(a)
    print(a)

def process_table_data(table_string):
    rows = table_string.strip().split('\n')
    table_data = [row.split() for row in rows]
    return table_data

def read_docx_with_order(file_path):
    doc = Document(file_path)
    content = []

    for para in doc.paragraphs:
        # Thêm văn bản vào danh sách
        if para.text.strip():
            content.append({"type": "text", "content": para.text})
        
        # Kiểm tra các phần tử trong run để tìm hình ảnh
        for run in para.runs:
            if "Graphic" in run._element.xml:
                for shape in run._element.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
                    image_rel_id = shape.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    image_part = doc.part.related_parts[image_rel_id]
                    image = Image.open(io.BytesIO(image_part.blob))
                    content.append({"type": "image", "content": image})

    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content.append({"type": "table", "content": table_data})

    return content

value = [{'type': 'text', 'content': '- Gồm một dòng duy nhất chứa hai số nguyên 𝑢 và 𝑣 (1 ≤ 𝑢, 𝑣 ≤2 × 109). '},
        {'type': 'text', 'content': 'Ví dụ: '},
        {'type': 'text', 'content': 'Ví dụ:'},
        {'type': 'text', 'content': 'Ví dụ:'},
        {'type': 'text', 'content': 'Ví dụ:'},
        ]

data = [{'INPUT': ['2 4', '3 7', '123 456'], 'OUTPUT': ['5.00', '14.50', '55766.25']},
{'INPUT': ['1 3 2 1 5 3 3 5', '3', '7 3 9 1 11 3 9 5', '12 4 13 2 15 1 14 4', '5 7 6 5 9 7 7 9'], 'OUTPUT': ['1', '', '', '', '']},
{'standard input': ['4', '10 30 40 20'], 'standard output': ['30', '']},
{'standard input': ['abhcdit', 'zhyixtw'], 'standard output': ['3', '']},
{'standard input': ['11', '3 4', '14', '3 2'], 'standard output': ['4 2', '6', '4 5', '8']}]

try:

    file_names = ["Bài1.docx","Bài2.docx","Bài3.docx","Bài4.docx","Bài5.docx"]
    cnt = 0
    for i in range(len(file_names)):
        arr = read_docx_with_order(file_names[i])
        title_unit = arr[0]['content']
        st.title(title_unit)
        content = arr[1:]
        id = content.index(value[i]) + 1
        content.insert(id,content[-1])
        content.pop()
        for item in content:
            if item["type"] == "text":
                st.write(item["content"])
            elif item["type"] == "image":
                st.image(item["content"])
            elif item["type"] == "table":
                df = pd.DataFrame(data[cnt])
                df.index = [''] * df.shape[0]
                st.table(df)
                cnt += 1
except Exception as e:
    st.error(f"Không thể đọc file: {e}")